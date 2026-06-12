from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 101, 119)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(22, 38, 58)
COPPER = RGBColor(181, 93, 44)
TEAL = RGBColor(37, 111, 110)
SAND = RGBColor(244, 233, 219)
SOFT = RGBColor(247, 242, 234)
LINE = RGBColor(214, 200, 182)
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
PHASE_FILLS = ["F7EFE6", "E8F1EF", "EEF1F5"]
PHASE_ACCENTS = [COPPER, TEAL, DARK_BLUE]

PHASES = [
    {
        "code": "FASE 1",
        "title": "Validering og Tech-Træning",
        "months": "Måned 1-6",
        "focus": "B2C-fokus med kendte pilotkunder, benchmark mod manuel proces og træning af AI-kæden fra lokalplan til mængder.",
        "north_star": "Gennemføre de første 3 pilotprojekter med kendte kunder i parløb med typehus-reference og løfte præcisionen nok til at gøre workflowet bank- og myndighedsklart.",
        "rows": [
            [
                "Måned 1-2",
                "Stabilisere data-pipelinen for DAR, MAT, BBR og lokalplaner. Indføre confidence scoring på dokumentfortolkning.",
                "Udvælge 3 kendte pilotkunder og definere benchmark, succesmål og betalingshypoteser.",
                "Etablere human-review SOP, audit trail og versionsstyring af alle regel- og mængdeoutput.",
                "100 dokumenter i træningskorpus og første adresseanalyse på minutter frem for dage.",
            ],
            [
                "Måned 3-4",
                "Automatisere første version af BIM-generering og mængdeberegning fra brief og inspirationsbilleder.",
                "Køre 2 aktive pilotforløb parallelt med manuel referenceproces.",
                "Etablere review-model med ekstern ingeniør-/forsikringspartner på kritiske sager.",
                "Lokalplanspræcision >85% og mængdeafvigelse <10%.",
            ],
            [
                "Måned 5-6",
                "Koble compliance-output direkte til myndighedspakke-skelet og stramme edge-case håndtering.",
                "Afslutte alle 3 pilots og omsætte læringer til produkt- og prisprioritering.",
                "Definere go/no-go model for hvornår en sag må fortsætte uden fuld manuel review.",
                "Lokalplanspræcision >95%, BIM-afvigelse <5% og første bygbar model på under 60 minutter.",
            ],
        ],
        "kpis": [
            "AI-præcision på lokalplansfortolkning >95% mod manuel benchmark",
            "BIM-mængdeafvigelse <5%",
            "Tre gennemførte pilotforløb med dokumenteret fejlmønster og læringslog",
        ],
    },
    {
        "code": "FASE 2",
        "title": "Det Kommercielle Gennembrud",
        "months": "Måned 7-12",
        "focus": "Første 100% ukendte kunde og den første reelle demonstration af bankability, myndighedsgang og entreprenøraccept gennem platformen.",
        "north_star": "Lande den første cold customer end-to-end og bevise, at ArchAI ikke kun kan modellere og beregne, men også bære et rigtigt købsscenarie i markedet.",
        "rows": [
            [
                "Måned 7-8",
                "Lancere betalte mikro-transaktioner for compliance, koncept, mængder og ansøgningspakker.",
                "Starte inbound og outbound test mod ukendte B2C-kunder og indsamle konverteringsdata.",
                "Bygge standardiseret bankpakke med budget, risikolog og dokumentpakke.",
                "Første betalende cold leads og de første funnel-data på unlock-strukturen.",
            ],
            [
                "Måned 9-10",
                "Produktisere eksportpakker til myndighedsproces og exception handling.",
                "Føre én cold customer gennem bankdialog, myndighed og entreprenørforhandling.",
                "Indføre anti-leakage vilkår og standardiseret udbudspakke til entreprenører.",
                "Mindst én bankdialog gennemført med ArchAI-materiale som primært beslutningsgrundlag.",
            ],
            [
                "Måned 11-12",
                "Stramme produktet til ét gentageligt golden path-flow og tydelig KPI-telemetri.",
                "Dokumentere første referencecase med reel betalingsvilje og partneraccept.",
                "Omsætte bruger- og partnerlæring til B2B-salgsmateriale og enterprise demo narrative.",
                "Positiv NPS og første færdigbyggede reference uden kritiske sager, når byggecyklussen lukker.",
            ],
        ],
        "kpis": [
            "Første cold customer gennem platformen fra brief til bankdialog",
            "Positiv NPS efter første afsluttede referenceforløb",
            "Mikro-transaktioner lanceret med målbare konverteringsrater",
        ],
    },
    {
        "code": "FASE 3",
        "title": "B2B Pivot og Skalering",
        "months": "Måned 13-18",
        "focus": "Pakke det lærte ind som moduler, white-label og API'er til typehusfirmaer, så softwareforretningen kan bære højere margin og lavere rådgiveransvar.",
        "north_star": "Lukke de første 2-3 enterprise aftaler og flytte forretningen fra assisteret casework til recurring software revenue.",
        "rows": [
            [
                "Måned 13-14",
                "Pakke kernefunktioner i moduler: compliance engine, konceptgenerator, BIM/mængder og myndighedspakke.",
                "Positionere ArchAI som white-label salgs- og projekteringsmotor over for mellemstore typehusaktører.",
                "Omtegne kontraktmodellen så partneren bærer endeligt rådgiveransvar i enterprise-setup.",
                "Demo-klar enterprise version og første design-partner pipeline.",
            ],
            [
                "Måned 15-16",
                "Bygge multi-tenant features, roller, account governance og API-adgang til udvalgte outputs.",
                "Køre strukturerede enterprise pilots med 2-3 konkrete kundeemner.",
                "Måle implementation time, cost-to-serve og ændringsønsker pr. enterprise account.",
                "Mindst 2 enterprise pilots med defineret kommerciel beslutningsdato.",
            ],
            [
                "Måned 17-18",
                "Productize onboarding, dokumentation og usage metering til licens- og API-prissætning.",
                "Lukke de første 2-3 aftaler med mellemstore danske typehusfirmaer.",
                "Flytte intern QA fra default review til exception review for at beskytte marginen.",
                "Første MRR, dokumenteret pipeline til de næste 12 måneder og elimineret internt rådgiveransvar som vækstbremse.",
            ],
        ],
        "kpis": [
            "Første faste månedlige abonnementsindtægter",
            "2-3 enterprise design partners eller kunder i aktiv implementation",
            "Review-arbejde flyttet mod exception-only på gentagelige sager",
        ],
    },
]

PROGRAM_RULES = [
    ("Proof before scale", "Volumen må ikke jagtes før bankability, compliance-præcision og myndighedspakke er bevist på rigtige cases."),
    ("Human review as exception", "Den manuelle ekspertrolle skal skubbes mod grænsetilfælde, ikke blive standardforløb på alle sager."),
    ("B2C as training loop", "B2C bruges som datamotor og læringsrum; B2B er den langsigtede margin- og distributionsmotor."),
    ("Liability boundary", "Ved enterprise-setup skal partneren bære det endelige rådgiveransvar, så ArchAI kan skalere som softwarevirksomhed."),
]


def load_table_geometry_module():
    script_path = Path(
        r"C:\Users\sidse\.codex\plugins\cache\openai-primary-runtime\documents\26.513.11550\skills\documents\scripts\table_geometry.py"
    )
    spec = importlib.util.spec_from_file_location("table_geometry", script_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry_module()


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D6C8B6", size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_paragraph_border_bottom(paragraph, color="D6C8B6", size=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def build_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ArchAI | Strategisk roadmap")
    set_run_font(run, size=9.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("Fortroligt | Maj 2026")
    set_run_font(run, size=9.5, color=MUTED, bold=False)


def add_text(doc_or_cell, text, *, style=None, size=11, color=BLACK, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0):
    paragraph = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_lead_callout(doc, title_text, body_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table_geometry.apply_table_geometry(
        table,
        [PAGE_WIDTH_DXA],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell, CALLOUT_FILL)
    set_cell_borders(cell, color="D6C8B6", size=8)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(4)
    run1 = p1.add_run(title_text)
    set_run_font(run1, size=11, color=MUTED, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    run2 = p2.add_run(body_text)
    set_run_font(run2, size=11, color=BLACK)
    add_text(doc, "", after=2)


def add_metadata_grid(doc):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry.apply_table_geometry(
        table,
        [4680, 4680],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    pairs = [
        ("Horisont", "18 måneder"),
        ("Primær mission", "Fra B2C-validering til B2B SaaS"),
        ("Dokumenttype", "Strategisk roadmap med milestones"),
        ("Målgruppe", "Founders, investorer og design partners"),
    ]
    for idx, pair in enumerate(pairs):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "FAF7F2")
        set_cell_borders(cell, color="D6C8B6", size=6)
        p_label = cell.paragraphs[0]
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(2)
        run_label = p_label.add_run(pair[0])
        set_run_font(run_label, size=10, color=MUTED, bold=True)
        p_value = cell.add_paragraph()
        p_value.paragraph_format.space_before = Pt(0)
        p_value.paragraph_format.space_after = Pt(0)
        run_value = p_value.add_run(pair[1])
        set_run_font(run_value, size=11, color=BLACK, bold=False)
    add_text(doc, "", after=2)


def add_overview_table(doc):
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [1560, 1560, 6240]
    table_geometry.apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    headers = ["Fase", "Måneder", "Primært udfald"]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        shade_cell(cell, TABLE_FILL)
        set_cell_borders(cell, color="C7D0DB", size=8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col < 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10.5, color=NAVY, bold=True)

    summaries = [
        ("Fase 1", "1-6", "Pilotprojekter, benchmark og AI-præcision >95% på lokalplansfortolkning."),
        ("Fase 2", "7-12", "Første cold customer og dokumenteret bankability/myndighedsgang."),
        ("Fase 3", "13-18", "White-label moduler, 2-3 enterprise aftaler og første MRR."),
    ]
    for row_idx, summary in enumerate(summaries, start=1):
        for col_idx, value in enumerate(summary):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell, color="D6C8B6", size=6)
            shade_cell(cell, "FFFDF9" if row_idx % 2 == 1 else "FAF7F2")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=10.5, color=BLACK, bold=(col_idx == 0))
    add_text(doc, "", after=2)


def add_phase_banner(doc, phase_index, phase):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry.apply_table_geometry(
        table,
        [2100, 7260],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(left, PHASE_FILLS[phase_index])
    shade_cell(right, "FFFDF8")
    set_cell_borders(left, color="D6C8B6", size=8)
    set_cell_borders(right, color="D6C8B6", size=8)

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_before = Pt(2)
    p_left.paragraph_format.space_after = Pt(0)
    run_code = p_left.add_run(phase["code"])
    set_run_font(run_code, size=11, color=PHASE_ACCENTS[phase_index], bold=True)
    p_left_2 = left.add_paragraph()
    p_left_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left_2.paragraph_format.space_before = Pt(0)
    p_left_2.paragraph_format.space_after = Pt(0)
    run_months = p_left_2.add_run(phase["months"])
    set_run_font(run_months, size=12, color=NAVY, bold=True)

    p_right = right.paragraphs[0]
    p_right.paragraph_format.space_before = Pt(0)
    p_right.paragraph_format.space_after = Pt(2)
    run_title = p_right.add_run(phase["title"])
    set_run_font(run_title, size=16, color=NAVY, bold=True)
    p_right_2 = right.add_paragraph()
    p_right_2.paragraph_format.space_before = Pt(0)
    p_right_2.paragraph_format.space_after = Pt(0)
    run_focus = p_right_2.add_run(phase["focus"])
    set_run_font(run_focus, size=10.5, color=BLACK)
    add_text(doc, "", after=2)


def add_phase_table(doc, phase_index, phase):
    table = doc.add_table(rows=1 + len(phase["rows"]), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [940, 1640, 1640, 1640, 3500]
    table_geometry.apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    headers = ["Tidsrum", "Teknologisk", "Kommercielt", "Operationelt", "Exit gate / KPI"]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        shade_cell(cell, TABLE_FILL)
        set_cell_borders(cell, color="C7D0DB", size=8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10, color=NAVY, bold=True)

    for row_idx, row in enumerate(phase["rows"], start=1):
        fill = "FFFDF9" if row_idx % 2 == 1 else "FAF7F2"
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(cell, fill)
            set_cell_borders(cell, color="D6C8B6", size=6)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=BLACK, bold=(col_idx == 0))
    add_text(doc, "", after=2)


def add_kpi_box(doc, phase_index, phase):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry.apply_table_geometry(
        table,
        [2500, 6860],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell, color="D6C8B6", size=8)
    shade_cell(left, PHASE_FILLS[phase_index])
    shade_cell(right, CALLOUT_FILL)

    p1 = left.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run("North star")
    set_run_font(run1, size=10.5, color=MUTED, bold=True)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(phase["north_star"])
    set_run_font(run2, size=10.5, color=BLACK, bold=False)

    rp1 = right.paragraphs[0]
    rp1.paragraph_format.space_before = Pt(0)
    rp1.paragraph_format.space_after = Pt(2)
    rr1 = rp1.add_run("Målepunkter")
    set_run_font(rr1, size=10.5, color=MUTED, bold=True)
    for idx, item in enumerate(phase["kpis"]):
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2 if idx < len(phase["kpis"]) - 1 else 0)
        run_label = p.add_run(f"KPI {idx + 1}: ")
        set_run_font(run_label, size=10.5, color=NAVY, bold=True)
        run_value = p.add_run(item)
        set_run_font(run_value, size=10.5, color=BLACK)
    add_text(doc, "", after=2)


def add_program_rules(doc):
    add_text(doc, "Cross-phase operating rules", style="Heading 2", size=13, color=BLUE, bold=True, after=7)
    table = doc.add_table(rows=len(PROGRAM_RULES), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry.apply_table_geometry(
        table,
        [2200, 7160],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    for row_idx, rule in enumerate(PROGRAM_RULES):
        left = table.cell(row_idx, 0)
        right = table.cell(row_idx, 1)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(left, "FAF7F2")
        shade_cell(right, "FFFDF9")
        set_cell_borders(left, color="D6C8B6", size=6)
        set_cell_borders(right, color="D6C8B6", size=6)
        p_left = left.paragraphs[0]
        p_left.paragraph_format.space_before = Pt(0)
        p_left.paragraph_format.space_after = Pt(0)
        run_left = p_left.add_run(rule[0])
        set_run_font(run_left, size=10.5, color=NAVY, bold=True)
        p_right = right.paragraphs[0]
        p_right.paragraph_format.space_before = Pt(0)
        p_right.paragraph_format.space_after = Pt(0)
        run_right = p_right.add_run(rule[1])
        set_run_font(run_right, size=10.5, color=BLACK)


def build_doc(output_path: Path):
    doc = Document()
    configure_document(doc)
    build_header_footer(doc.sections[0])

    top_meta = add_text(doc, "ArchAI", size=12, color=MUTED, bold=True, after=4)
    top_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = add_text(doc, "Strategisk Roadmap og Milestones", style="Title", size=24, color=NAVY, bold=True, after=4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = add_text(
        doc,
        "18 måneder fra B2C-validering til B2B SaaS-skalering",
        size=13.5,
        color=MUTED,
        after=10,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph_border_bottom(subtitle)
    add_text(doc, "", after=4)

    add_metadata_grid(doc)
    add_lead_callout(
        doc,
        "Program thesis",
        "ArchAI skal først bevise, at AI-kæden kan levere bankable, compliance-robuste og dokumentérbare beslutningspakker i rigtige B2C-forløb. Derefter skal de samme workflows pakkes ind som white-label moduler og API'er til typehusfirmaer, hvor marginen og distributionskraften er stærkere.",
    )
    add_overview_table(doc)

    add_text(doc, "Programlogik", style="Heading 2", size=13, color=BLUE, bold=True, after=7)
    add_text(
        doc,
        "Roadmappet er bygget baglæns fra den værdi, en investor eller design partner faktisk skal tro på: ikke kun at ArchAI kan generere flotte output, men at platformen kan bære reelle købsscenarier, myndighedsgange og enterprise-adoption med en forsvarlig liability boundary.",
        size=11,
        color=BLACK,
        after=8,
    )

    for idx, phase in enumerate(PHASES):
        doc.add_page_break()
        add_phase_banner(doc, idx, phase)
        add_phase_table(doc, idx, phase)
        add_kpi_box(doc, idx, phase)

    add_text(doc, "", after=4)
    add_program_rules(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", required=True)
    args = parser.parse_args()
    build_doc(Path(args.out))


if __name__ == "__main__":
    main()
